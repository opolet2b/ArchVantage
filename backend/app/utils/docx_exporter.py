import os
import zipfile
import shutil
import tempfile
import uuid
import re
from docx import Document
from docx.shared import Inches

def add_markdown_text(doc, text):
    """Parses basic markdown (**, *, \n) and appends to the docx Document natively."""
    if not text:
        return
    paragraphs = text.split('\n')
    for line in paragraphs:
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        bold_parts = re.split(r'(\*\*.*?\*\*)', line)
        for b_part in bold_parts:
            if b_part.startswith('**') and b_part.endswith('**'):
                p.add_run(b_part[2:-2]).bold = True
            else:
                italic_parts = re.split(r'(\*.*?\*)', b_part)
                for i_part in italic_parts:
                    if i_part.startswith('*') and i_part.endswith('*'):
                        p.add_run(i_part[1:-1]).italic = True
                    else:
                        p.add_run(i_part)

def generate_memo_docx(markdown_content: str) -> str:
    temp_dir = tempfile.mkdtemp()
    doc = Document()
    doc.add_heading("Architecture Memo", 0)
    
    paragraphs = markdown_content.split('\n')
    table_lines = []
    
    def process_table():
        if not table_lines:
            return
        
        # Filter out the separator line (e.g. |---|---|)
        data_lines = [line for line in table_lines if not re.match(r'^[\s|:\-]+$', line)]
        if not data_lines:
            table_lines.clear()
            return
            
        # Determine columns from first data line
        cols = [c.strip() for c in data_lines[0].strip('|').split('|')]
        table = doc.add_table(rows=len(data_lines), cols=len(cols))
        table.style = 'Table Grid'
        
        for r_idx, line in enumerate(data_lines):
            cells = [c.strip() for c in line.strip('|').split('|')]
            for c_idx, cell_text in enumerate(cells):
                if c_idx < len(table.columns):
                    p = table.rows[r_idx].cells[c_idx].paragraphs[0]
                    _apply_inline_markdown(p, cell_text)
                    if r_idx == 0:
                        # Make header bold
                        for run in p.runs:
                            run.bold = True
                            
        table_lines.clear()

    for line in paragraphs:
        line = line.strip()
        
        # Check if line is part of a markdown table
        if line.startswith('|') and line.endswith('|'):
            table_lines.append(line)
            continue
        else:
            process_table()
            
        if not line:
            continue
        
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            text = line[2:]
            _apply_inline_markdown(p, text)
        else:
            p = doc.add_paragraph()
            _apply_inline_markdown(p, line)
            
    # Process any table left at the end
    process_table()
            
    final_docx = os.path.join(tempfile.gettempdir(), f"memo_export_{uuid.uuid4().hex}.docx")
    doc.save(final_docx)
    return final_docx

def _apply_inline_markdown(p, text):
    """Helper to apply bold/italic to a paragraph"""
    bold_parts = re.split(r'(\*\*.*?\*\*)', text)
    for b_part in bold_parts:
        if b_part.startswith('**') and b_part.endswith('**'):
            p.add_run(b_part[2:-2]).bold = True
        else:
            italic_parts = re.split(r'(\*.*?\*)', b_part)
            for i_part in italic_parts:
                if i_part.startswith('*') and i_part.endswith('*'):
                    p.add_run(i_part[1:-1]).italic = True
                else:
                    p.add_run(i_part)

def generate_scenario_docx(payload: dict) -> str:
    temp_dir = tempfile.mkdtemp()
    
    dummy_png_path = os.path.join(temp_dir, "dummy.png")
    with open(dummy_png_path, "wb") as f:
        f.write(b'\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82')

    doc = Document()
    doc.add_heading("Architectural Scenario Impact Analysis", 0)

    doc.add_heading("1. Scenario Configuration", level=1)
    doc.add_paragraph(f"Action: {payload.get('action', 'none')}", style='List Bullet')
    doc.add_paragraph(f"Target Technology: {payload.get('targetTech', 'N/A')}", style='List Bullet')
    doc.add_paragraph(f"Custom Prompt: {payload.get('customPrompt', 'N/A')}", style='List Bullet')

    doc.add_heading("2. Structural Risk Assessment", level=1)
    doc.add_paragraph(f"Risk Score: {payload.get('risk_score', 0)} / 100")
    doc.add_paragraph("Rationale:")
    add_markdown_text(doc, payload.get('risk_rationale', ''))

    doc.add_heading("3. Architectural Diagrams", level=1)

    images_to_replace = []
    
    if payload.get("baseline_svg"):
        doc.add_heading("Baseline (As-Is)", level=2)
        pic = doc.add_picture(dummy_png_path, width=Inches(6.0))
        images_to_replace.append(payload["baseline_svg"])

    if payload.get("tobe_svg"):
        doc.add_heading("To-Be (Simulation)", level=2)
        pic = doc.add_picture(dummy_png_path, width=Inches(6.0))
        images_to_replace.append(payload["tobe_svg"])

    doc.add_heading("4. TOGAF ADM Phase Impacts", level=1)
    phases = payload.get("phases", [])
    for p in phases:
        doc.add_heading(f"Phase {p['phase']}: {p['name']}", level=2)
        doc.add_paragraph(f"Risk Level: {p['impact']}")
        add_markdown_text(doc, p['desc'])

    temp_docx = os.path.join(temp_dir, "temp.docx")
    doc.save(temp_docx)

    extract_dir = os.path.join(temp_dir, "extracted")
    os.makedirs(extract_dir, exist_ok=True)
    with zipfile.ZipFile(temp_docx, 'r') as zip_ref:
        zip_ref.extractall(extract_dir)

    media_dir = os.path.join(extract_dir, "word", "media")
    if os.path.exists(media_dir):
        image_files = sorted([f for f in os.listdir(media_dir) if f.endswith(".png")])
        for i, png_file in enumerate(image_files):
            if i < len(images_to_replace):
                svg_content = images_to_replace[i]
                base_name = os.path.splitext(png_file)[0]
                svg_file = base_name + ".svg"
                
                with open(os.path.join(media_dir, svg_file), "w", encoding="utf-8") as f:
                    f.write(svg_content)
                
                os.remove(os.path.join(media_dir, png_file))

                rels_file = os.path.join(extract_dir, "word", "_rels", "document.xml.rels")
                with open(rels_file, "r", encoding="utf-8") as f:
                    rels_content = f.read()
                rels_content = rels_content.replace(png_file, svg_file)
                with open(rels_file, "w", encoding="utf-8") as f:
                    f.write(rels_content)

    ct_file = os.path.join(extract_dir, "[Content_Types].xml")
    with open(ct_file, "r", encoding="utf-8") as f:
        ct_content = f.read()
    if 'Extension="svg"' not in ct_content:
        insert_str = '<Default Extension="svg" ContentType="image/svg+xml"/>'
        ct_content = ct_content.replace('</Types>', f'{insert_str}</Types>')
    with open(ct_file, "w", encoding="utf-8") as f:
        f.write(ct_content)

    final_docx = os.path.join(tempfile.gettempdir(), f"export_{uuid.uuid4().hex}.docx")
    with zipfile.ZipFile(final_docx, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, dirs, files in os.walk(extract_dir):
            for file in files:
                file_path = os.path.join(root, file)
                arcname = os.path.relpath(file_path, extract_dir)
                zipf.write(file_path, arcname)

    shutil.rmtree(temp_dir)
    return final_docx

def generate_time_matrix_svg_string(apps: list) -> str:
    width = 600
    height = 600
    margin = 50
    chart_w = width - 2*margin
    chart_h = height - 2*margin

    colors = {
        "INVEST": "#22c55e", 
        "MIGRATE": "#f59e0b", 
        "ELIMINATE": "#ef4444", 
        "TOLERATE": "#64748b" 
    }

    svg = f'<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 {width} {height}" width="{width}" height="{height}">'
    svg += f'<rect x="{margin}" y="{margin}" width="{chart_w}" height="{chart_h}" fill="#f8fafc" stroke="#cbd5e1" stroke-width="2" />'

    mid_x = margin + chart_w / 2
    mid_y = margin + chart_h / 2
    svg += f'<line x1="{margin}" y1="{mid_y}" x2="{width-margin}" y2="{mid_y}" stroke="#94a3b8" stroke-width="2" stroke-dasharray="6,6" />'
    svg += f'<line x1="{mid_x}" y1="{margin}" x2="{mid_x}" y2="{height-margin}" stroke="#94a3b8" stroke-width="2" stroke-dasharray="6,6" />'

    style = "font-family: Arial, sans-serif; font-size: 20px; font-weight: bold; fill: #cbd5e1;"
    svg += f'<text x="{margin + 20}" y="{margin + 30}" style="{style}">MIGRATE</text>'
    svg += f'<text x="{width - margin - 90}" y="{margin + 30}" style="{style}">INVEST</text>'
    svg += f'<text x="{margin + 20}" y="{height - margin - 20}" style="{style}">ELIMINATE</text>'
    svg += f'<text x="{width - margin - 120}" y="{height - margin - 20}" style="{style}">TOLERATE</text>'
    
    axis_style = "font-family: Arial, sans-serif; font-size: 14px; font-weight: bold; fill: #64748b; text-anchor: middle;"
    svg += f'<text x="{mid_x}" y="{height - margin + 30}" style="{axis_style}">Technical Health (0-10)</text>'
    svg += f'<text x="{-mid_y}" y="{margin - 30}" style="{axis_style}" transform="rotate(-90)">Business Value (0-10)</text>'

    placed_points = []
    import math
    for i, app in enumerate(apps):
        th = float(app.get("technicalHealth", 0))
        bv = float(app.get("businessValue", 0))
        cost = float(app.get("runCost") or 10000)
        quadrant = str(app.get("quadrant", "TOLERATE")).upper()

        cx = margin + (th / 10.0) * chart_w
        cy = margin + chart_h - (bv / 10.0) * chart_h
        
        r = max(10.0, min(30.0, cost / 4000.0)) 
        
        collision = True
        attempts = 0
        while collision and attempts < 30:
            collision = False
            for px, py, pr in placed_points:
                dist = ((cx - px)**2 + (cy - py)**2)**0.5
                if dist < (r + pr + 2):
                    angle = attempts * (math.pi / 4)
                    shift = 5 + attempts
                    cx += math.cos(angle) * shift
                    cy += math.sin(angle) * shift
                    collision = True
                    break
            attempts += 1
            
        placed_points.append((cx, cy, r))

        color = colors.get(quadrant, "#3b82f6")
        
        svg += f'<circle cx="{cx}" cy="{cy}" r="{r}" fill="{color}" fill-opacity="0.8" stroke="#ffffff" stroke-width="2" />'
        label_style = "font-family: Arial, sans-serif; font-size: 14px; font-weight: bold; fill: #ffffff; text-anchor: middle;"
        svg += f'<text x="{cx}" y="{cy + 5}" style="{label_style}">{i + 1}</text>'

    svg += '</svg>'
    return svg


def generate_time_matrix_docx(payload: dict) -> str:
    temp_dir = tempfile.mkdtemp()
    apps = payload.get("apps", [])
    
    dummy_png_path = os.path.join(temp_dir, "dummy.png")
    with open(dummy_png_path, "wb") as f:
        f.write(b'\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82')

    doc = Document()
    doc.add_heading("Application Portfolio: TIME Matrix Analysis", 0)

    doc.add_heading("1. Matrix Overview", level=1)
    
    svg_content = generate_time_matrix_svg_string(apps)
    doc.add_picture(dummy_png_path, width=Inches(6.0))
    
    doc.add_heading("2. Application Portfolio", level=1)
    
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '#'
    hdr_cells[1].text = 'Application'
    hdr_cells[2].text = 'Quadrant'
    hdr_cells[3].text = 'Tech Health'
    hdr_cells[4].text = 'Bus Value'
    hdr_cells[5].text = 'Run Cost'
    hdr_cells[6].text = 'Risk Profile'
    
    for cell in hdr_cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    for i, app in enumerate(apps):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = str(app.get("name", ""))
        row_cells[2].text = str(app.get("quadrant", ""))
        row_cells[3].text = str(app.get("technicalHealth", ""))
        row_cells[4].text = str(app.get("businessValue", ""))
        cost = app.get("runCost")
        row_cells[5].text = f"${cost:,}" if isinstance(cost, (int, float)) else str(cost or "N/A")
        row_cells[6].text = str(app.get("riskProfile", ""))

    doc.add_heading("3. Evidence and Citations", level=1)
    for app in apps:
        citations = app.get("citations", [])
        if citations:
            doc.add_heading(f"{app.get('name')} ({app.get('quadrant')})", level=2)
            for cit in citations:
                doc.add_paragraph(str(cit), style='List Bullet')

    temp_docx = os.path.join(temp_dir, "temp.docx")
    doc.save(temp_docx)

    extract_dir = os.path.join(temp_dir, "extracted")
    os.makedirs(extract_dir, exist_ok=True)
    with zipfile.ZipFile(temp_docx, 'r') as zip_ref:
        zip_ref.extractall(extract_dir)

    media_dir = os.path.join(extract_dir, "word", "media")
    if os.path.exists(media_dir):
        image_files = sorted([f for f in os.listdir(media_dir) if f.endswith(".png")])
        if image_files:
            png_file = image_files[0]
            base_name = os.path.splitext(png_file)[0]
            svg_file = base_name + ".svg"
            
            with open(os.path.join(media_dir, svg_file), "w", encoding="utf-8") as f:
                f.write(svg_content)
            
            os.remove(os.path.join(media_dir, png_file))

            rels_file = os.path.join(extract_dir, "word", "_rels", "document.xml.rels")
            with open(rels_file, "r", encoding="utf-8") as f:
                rels_content = f.read()
            rels_content = rels_content.replace(png_file, svg_file)
            with open(rels_file, "w", encoding="utf-8") as f:
                f.write(rels_content)

    ct_file = os.path.join(extract_dir, "[Content_Types].xml")
    with open(ct_file, "r", encoding="utf-8") as f:
        ct_content = f.read()
    if 'Extension="svg"' not in ct_content:
        insert_str = '<Default Extension="svg" ContentType="image/svg+xml"/>'
        ct_content = ct_content.replace('</Types>', f'{insert_str}</Types>')
    with open(ct_file, "w", encoding="utf-8") as f:
        f.write(ct_content)

    final_docx = os.path.join(tempfile.gettempdir(), f"time_matrix_export_{uuid.uuid4().hex}.docx")
    with zipfile.ZipFile(final_docx, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, dirs, files in os.walk(extract_dir):
            for file in files:
                file_path = os.path.join(root, file)
                arcname = os.path.relpath(file_path, extract_dir)
                zipf.write(file_path, arcname)

    shutil.rmtree(temp_dir)
    return final_docx
