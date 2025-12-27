import os
import shutil
from typing import List, Optional
import chromadb
from llama_index.core import VectorStoreIndex, SimpleDirectoryReader, StorageContext, Settings, Document
from llama_index.vector_stores.chroma import ChromaVectorStore
from llama_index.embeddings.ollama import OllamaEmbedding
from llama_index.core.node_parser import SentenceSplitter

class RAGService:
    def __init__(self):
        self.persist_directory = "./chroma_db"
        self._initialized = False
        
        try:
            # Configure Settings
            # Using Ollama Embeddings to avoid local CPU hangs and leverage GPU if available
            print("[RAGService] Connecting to Ollama for Embeddings (model: nomic-embed-text)...")
            Settings.embed_model = OllamaEmbedding(
                model_name="nomic-embed-text",
                base_url="http://localhost:11434",
                ollama_additional_kwargs={"mirostat": 0}
            )
            Settings.text_splitter = SentenceSplitter(chunk_size=1000, chunk_overlap=200)
            
            # Initialize Chroma Client
            # Use PersistentClient for data retention
            print(f"[RAGService] Initializing PersistentClient at {self.persist_directory}")
            self.chroma_client = chromadb.PersistentClient(path=self.persist_directory)
            # Using v2 collection to support 768-dim embeddings (Ollama/Nomic) instead of old 384-dim
            self.chroma_collection = self.chroma_client.get_or_create_collection("chatbot_rag_v2")
            
            # Set up Vector Store and Storage Context
            self.vector_store = ChromaVectorStore(chroma_collection=self.chroma_collection)
            self.storage_context = StorageContext.from_defaults(vector_store=self.vector_store)
            
            # Load index from storage if it exists, otherwise create empty
            try:
                self.index = VectorStoreIndex.from_vector_store(
                    self.vector_store,
                    storage_context=self.storage_context
                )
            except Exception:
                self.index = VectorStoreIndex.from_documents(
                    [], storage_context=self.storage_context
                )
            
            self._initialized = True
        except Exception as e:
            print(f"RAGService initialization failed: {e}")
            print("RAG features will be disabled.")
            self.chroma_client = None
            self.chroma_collection = None
            self.vector_store = None
            self.storage_context = None
            self.index = None

    def ingest_file(self, file_path: str, conversation_id: Optional[str] = None, metadata: Optional[dict] = None):
        print(f"[RAGService] Starting ingestion for: {file_path}")
        try:
            # LlamaIndex SimpleDirectoryReader handles various file types automatically
            # We explicitly check for OLE (binary .doc) vs ZIP (.docx) to avoid reading binary as text
            
            # Check magic bytes
            with open(file_path, 'rb') as f:
                header = f.read(8)
            
            is_ole = header.startswith(b'\xd0\xcf\x11\xe0')
            is_zip = header.startswith(b'PK\x03\x04')
            
            if is_ole:
                print(f"[RAGService] Detected binary OLE file (legacy .doc): {file_path}")
                # MarkItDown *might* handle it if it supports .doc, or we need another tool.
                # If MarkItDown supports it via 'unstructured' or similar:
                try:
                    from markitdown import MarkItDown
                    md = MarkItDown()
                    result = md.convert(file_path)
                    text = result.text_content
                    documents = [Document(text=text)]
                except Exception as e:
                    print(f"[RAGService] MarkItDown failed on .doc file: {e}")
                    # Fallback or error message for user
                    return {"status": "error", "error": "Legacy .doc format detected. Please save as .docx and try again."}
                    
            elif is_zip or file_path.lower().endswith('.docx'):
                 # It's likely a valid .docx
                from markitdown import MarkItDown
                md = MarkItDown()
                result = md.convert(file_path)
                text = result.text_content
                
                # Check for images to warn user
                try:
                    import docx
                    doc = docx.Document(file_path)
                    has_images = False
                    
                    # Check inline shapes (common for images)
                    if len(doc.inline_shapes) > 0:
                         has_images = True
                    else:
                        # Check relationships for any image types (header/footer images etc)
                        for rel in doc.part.rels.values():
                            if "image" in rel.reltype:
                                has_images = True
                                break
                                
                    if has_images:
                        warning_msg = (
                            "> [!WARNING]\n"
                            "> **Images Detected**: This document contains images or embedded objects which cannot be displayed here.\n"
                            "> To view this document with full visual fidelity, please **export it as a PDF** and import the PDF instead.\n\n"
                        )
                        text = warning_msg + text
                except Exception as e:
                    print(f"[RAGService] Failed to check for images in docx: {e}")
                
                documents = [Document(text=text)]
                
            else:
                # Default for other types
                documents = SimpleDirectoryReader(input_files=[file_path]).load_data()
                
            print(f"[RAGService] Loaded {len(documents)} document fragments from file.")
            
            if documents:
                # Add metadata
                for doc in documents:
                    if conversation_id:
                        doc.metadata["conversation_id"] = conversation_id
                    
                    doc.metadata["source"] = file_path
                    
                    # Add metadata
                for doc in documents:
                    if conversation_id:
                        doc.metadata["conversation_id"] = conversation_id
                    
                    doc.metadata["source"] = file_path
                    
                    # Add custom metadata (e.g. canvas_id)
                    if metadata:
                        for key, value in metadata.items():
                            doc.metadata[key] = value

                    # Excluded metadata keys to avoid errors with some vector stores if needed
                    doc.excluded_llm_metadata_keys = ["conversation_id", "source"]
                    if metadata:
                        doc.excluded_llm_metadata_keys.extend(metadata.keys())

                    doc.excluded_embed_metadata_keys = ["conversation_id", "source"]
                    if metadata:
                         doc.excluded_embed_metadata_keys.extend(metadata.keys())

                # Optimization: Split into nodes first
                print(f"[RAGService] Splitting {len(documents)} documents into nodes...")
                nodes = Settings.text_splitter.get_nodes_from_documents(documents)
                print(f"[RAGService] Created {len(nodes)} nodes.")
                
                # Debug: Test embedding generation explicitly
                if nodes:
                    print("[RAGService] DEBUG: Generating embedding for first node to verify model...")
                    import time
                    t0 = time.time()
                    # We can manually get embedding to show it works
                    _ = Settings.embed_model.get_text_embedding(nodes[0].get_content())
                    print(f"[RAGService] DEBUG: Embedding generation successful. Took {time.time()-t0:.2f}s")

                print(f"[RAGService] Inserting {len(nodes)} nodes one-by-one to trace progress...")
                for i, node in enumerate(nodes):
                     print(f"[RAGService] Inserting node {i+1}/{len(nodes)}...")
                     self.index.insert_nodes([node])
                     print(f"[RAGService] Node {i+1} inserted.")
                
                # Calculate total text length for heuristic checks
                total_text_len = sum(len(node.get_content()) for node in nodes)
                
                print(f"[RAGService] Vectorization complete. Ingested {len(nodes)} fragments. Total Chars: {total_text_len}")
                
                # Combine all text for frontend display
                full_text = "\n\n".join([node.get_content() for node in nodes])
                
                return {
                    "status": "success", 
                    "count": len(nodes), 
                    "text_length": total_text_len,
                    "doc_count": len(documents),
                    "full_text": full_text
                }
            
            print(f"[RAGService] WARNING: No content extracted from file: {file_path}")
            return {"status": "no_documents_found"}
        except Exception as e:
            print(f"Error ingesting file {file_path}: {e}")
            raise e

    MAX_TEXT_LENGTH = 10000

    def ingest_slideshow(self, file_path: str, conversation_id: Optional[str] = None, metadata: Optional[dict] = None, progress_callback=None):
        """
        Ingest a PowerPoint file using its pre-extracted JSON structure.
        Delegates to the specialized SlideshowIngestor.
        """
        try:
            from app.services.rag.slideshow_ingestor import slideshow_ingestor
            
            return slideshow_ingestor.ingest_slideshow(
                file_path=file_path,
                index=self.index,
                storage_context=self.storage_context,
                conversation_id=conversation_id,
                metadata=metadata,
                progress_callback=progress_callback
            )
        except Exception as e:
            print(f"[RAGService] Error ingesting slideshow: {e}")
            raise e

    def ingest_text(self, text: str, metadata: Optional[dict] = None):
        """
        Ingest raw text directly into the index.
        Useful for image descriptions or other generated content.
        """
        try:
            # Create a Document object
            doc = Document(text=text, metadata=metadata or {})
            
            # Split into nodes
            nodes = Settings.text_splitter.get_nodes_from_documents([doc])
            
            if nodes:
               print(f"[RAGService] Ingesting {len(nodes)} nodes from text content.")
               self.index.insert_nodes(nodes)
               return {"status": "success", "count": len(nodes)}
            
            return {"status": "no_content"}
        except Exception as e:
            print(f"[RAGService] Error ingesting text: {e}")
            raise e

    def ingest_folder(self, folder_path: str, chunk_size: int = 1000, chunk_overlap: int = 200, metadata: Optional[dict] = None):
        # Not strictly needed for the current flow but good to keep
        try:
            # Configure splitting dynamically
            Settings.text_splitter = SentenceSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
            
            documents = SimpleDirectoryReader(input_dir=folder_path, recursive=True).load_data()
            if documents:
                for doc in documents:
                    # Apply metadata if provided
                    if metadata:
                        for key, value in metadata.items():
                            doc.metadata[key] = value
                    pass 
                
                # Delete existing documents from index to avoid duplicates if re-ingesting?
                # For now, just insert (LlamaIndex might handle dupes or we accept them)
                # Actually, clearing legacy "data" ingestion might be good, but risky if we mix with uploads.
                # Let's keep append behavior for now.
                
                for doc in documents:
                    self.index.insert(doc)
                return {"status": "success", "count": len(documents)}
            return {"status": "no_documents_found"}
        except Exception as e:
            print(f"Error ingesting folder {folder_path}: {e}")
            return {"status": "error", "detail": str(e)}

    def query(self, query_text: str, k: int = 4, conversation_id: Optional[str] = None, filters: Optional[dict] = None):
        from llama_index.core.vector_stores import MetadataFilters, ExactMatchFilter
        
        # Build filters
        metadata_filters = None
        filter_list = []
        
        # Legacy support for conversation_id arg
        if conversation_id:
            filter_list.append(ExactMatchFilter(key="conversation_id", value=conversation_id))
            
        # Generic filters
        if filters:
            for key, value in filters.items():
                filter_list.append(ExactMatchFilter(key=key, value=value))
                
        if filter_list:
            metadata_filters = MetadataFilters(filters=filter_list)

        # Create retriever
        retriever = self.index.as_retriever(similarity_top_k=k, filters=metadata_filters)
        nodes = retriever.retrieve(query_text)
        
        return [node.get_content() for node in nodes]

    def search(
        self,
        query: str,
        conversation_id: Optional[str] = None,
        filters: Optional[dict] = None,
        k: int = 10
    ) -> List[dict]:
        """
        Search for relevant documents with metadata and scores.
        
        Unlike query() which returns only text content, this method
        returns structured results including metadata and relevance scores.
        
        Args:
            query: The search query text.
            conversation_id: Optional filter by conversation/collection ID.
            filters: Optional dictionary of metadata filters (e.g. {"canvas_id": "123"})
            k: Number of results to return.
            
        Returns:
            List of dicts with 'text', 'metadata', and 'score' keys.
        """
        try:
            print(f"[RAGService] Search requested. Query: '{query}' Filters: {filters}")
            if not self._initialized or self.index is None:
                print(f"[RAGService] Index not initialized. Returning empty.")
                return []
            
            print(f"[RAGService] Building filters...")
            from llama_index.core.vector_stores import MetadataFilters, ExactMatchFilter
            
             # Build filters
            metadata_filters = None
            filter_list = []
            
            # Legacy support for conversation_id arg
            if conversation_id:
                filter_list.append(ExactMatchFilter(key="conversation_id", value=conversation_id))
                
            # Generic filters
            if filters:
                for key, value in filters.items():
                    filter_list.append(ExactMatchFilter(key=key, value=value))
                    
            if filter_list:
                metadata_filters = MetadataFilters(filters=filter_list)
            
            # Create retriever
            print(f"[RAGService] Creating retriever (k={k})...")
            retriever = self.index.as_retriever(similarity_top_k=k, filters=metadata_filters)
            
            print(f"[RAGService] Executing retrieve('{query}')...")
            nodes = retriever.retrieve(query)
            print(f"[RAGService] Retrieved {len(nodes)} nodes.")
            
            results = []
            for node in nodes:
                results.append({
                    "text": node.get_content(),
                    "metadata": node.metadata if hasattr(node, 'metadata') else {},
                    "score": node.score if hasattr(node, 'score') else 0.0
                })
            
            return results
        except Exception as e:
            print(f"RAG search error (returning empty): {e}")
            return []

    def delete_conversation_embeddings(self, conversation_id: str):
        try:
            # Delete directly from Chroma collection
            self.chroma_collection.delete(where={"conversation_id": conversation_id})
            return True
        except Exception as e:
            print(f"Error deleting embeddings for conversation {conversation_id}: {e}")
            return False

    def list_documents(self, conversation_id: str) -> List[str]:
        upload_dir = f"data/uploads/{conversation_id}"
        if not os.path.exists(upload_dir):
            return []
        return os.listdir(upload_dir)

    def get_document_content(self, conversation_id: str, filename: str) -> str:
        file_path = f"data/uploads/{conversation_id}/{filename}"
        if not os.path.exists(file_path):
            return "File not found"
        
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            return f"Error reading file: {str(e)}"

    def delete_document(self, conversation_id: str, filename: str):
        file_path = f"data/uploads/{conversation_id}/{filename}"
        if os.path.exists(file_path):
            os.remove(file_path)
            
        try:
            # Delete from Chroma collection
            self.chroma_collection.delete(
                where={
                    "$and": [
                        {"conversation_id": conversation_id},
                        {"source": file_path}
                    ]
                }
            )
            return True
        except Exception as e:
            print(f"Error deleting embeddings for file {filename}: {e}")
            return False

    def reset_db(self):
        try:
            self.chroma_client.delete_collection("chatbot_rag")
            self.chroma_collection = self.chroma_client.create_collection("chatbot_rag")
            
            # Re-initialize index
            self.vector_store = ChromaVectorStore(chroma_collection=self.chroma_collection)
            self.storage_context = StorageContext.from_defaults(vector_store=self.vector_store)
            self.index = VectorStoreIndex.from_vector_store(
                self.vector_store,
                storage_context=self.storage_context
            )
        except Exception as e:
            print(f"Error resetting DB: {e}")

rag_service = RAGService()
