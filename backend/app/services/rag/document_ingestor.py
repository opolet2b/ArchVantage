import os
from typing import List, Optional, Dict, Any
from llama_index.core import Document, VectorStoreIndex, StorageContext
from llama_index.core import SimpleDirectoryReader

class DocumentIngestor:
    """
    Ingestor for standard documents (PDF, Docx, Text, etc).
    Supports progress reporting.
    """
    
    def __init__(self):
        pass

    def ingest_document(self, file_path: str, index: VectorStoreIndex, storage_context: StorageContext, conversation_id: Optional[str] = None, metadata: Optional[dict] = None, progress_callback=None, llm=None):
        """
        Ingest a document file.
        Supports .docx specific handling via MarkItDown.
        Args:
            llm: Optional LlamaIndex LLM instance for metadata extraction.
        """
        print(f"[DocumentIngestor] Starting ingestion for: {file_path}")
        
        try:
            structured_data = None
            is_docx = file_path.lower().endswith('.docx')
            
            # Word Document Handling
            if is_docx:
                from markitdown import MarkItDown
                import re
                
                md = MarkItDown()
                result = md.convert(file_path)
                text = result.text_content
                
                # Clean up base64 images from markdown to prevent context overflow
                text = re.sub(
                    r'!\[(.*?)\]\(data:image/[^)]+\)', 
                    lambda m: f"[Image: {m.group(1)}]" if m.group(1) else "[Image]", 
                    text
                )
                
                # Check for images to warn user
                try:
                    import docx
                    doc = docx.Document(file_path)
                    has_images = False
                    if len(doc.inline_shapes) > 0:
                        has_images = True
                    else:
                        for rel in doc.part.rels.values():
                            if "image" in rel.reltype:
                                has_images = True
                                break
                    if has_images:
                        warning_msg = (
                            "> [!WARNING]\n"
                            "> **Images Detected**: This document contains images or embedded objects which cannot be displayed here.\n"
                            "> To view this document with full visual fidelity, please **export it as a PDF** and import the PDF instead.\n\n"
                        )
                        text = warning_msg + text
                except Exception as e:
                    print(f"[DocumentIngestor] Failed to check for images in docx: {e}")
                
                documents = [Document(text=text)]
                
            else:
                # Check for Excel/CSV for structured extraction
                is_excel = file_path.lower().endswith(('.xlsx', '.xls'))
                is_csv = file_path.lower().endswith('.csv')

                if is_excel or is_csv:
                    try:
                        import pandas as pd
                        print(f"[DocumentIngestor] Extracting structured data from {'Excel' if is_excel else 'CSV'}...")
                        
                        if is_excel:
                            # Read all sheets
                            excel_dict = pd.read_excel(file_path, sheet_name=None)
                            # Combine all sheets into one DataFrame for structured data
                            all_dfs = []
                            text_parts = []
                            for sheet_name, sheet_df in excel_dict.items():
                                if not sheet_df.empty:
                                    all_dfs.append(sheet_df)
                                    # Fallback for to_markdown
                                    try:
                                        sheet_md = sheet_df.to_markdown(index=False)
                                    except:
                                        sheet_md = sheet_df.to_string(index=False)
                                    text_parts.append(f"### Sheet: {sheet_name}\n" + sheet_md)
                            
                            if all_dfs:
                                df = pd.concat(all_dfs, ignore_index=True)
                                text = "\n\n".join(text_parts)
                            else:
                                df = pd.DataFrame()
                                text = ""
                        else:
                            df = pd.read_csv(file_path)
                            try:
                                text = df.to_markdown(index=False)
                            except:
                                text = df.to_string(index=False)
                        
                        if df.empty:
                            print(f"[DocumentIngestor] Warning: {file_path} is empty.")
                            structured_data = {"columns": [], "rows": []}
                            documents = [Document(text="Empty table content")]
                        else:
                            # Convert to rows/columns
                            cols = df.columns.tolist()
                            raw_rows = df.values.tolist()
                            
                            # Sanitize values
                            def sanitize_val(v):
                                import math
                                try:
                                    if v is None: return None
                                    if isinstance(v, float) and math.isnan(v): return None
                                    if isinstance(v, (list, dict)): return str(v)
                                    return v
                                except: return str(v)

                            clean_rows = []
                            header = [str(c) for c in cols]
                            clean_rows.append(header)
                            
                            for r in raw_rows:
                                clean_rows.append([sanitize_val(v) for v in r])
                            
                            structured_data = {
                                "columns": header,
                                "rows": clean_rows
                            }
                            
                            documents = [Document(text=text)]
                            print(f"[DocumentIngestor] Structured extraction successful: {len(clean_rows)} rows found.")
                        
                    except Exception as e:
                        print(f"[DocumentIngestor] Failed structured extraction for {file_path}: {e}")
                        documents = SimpleDirectoryReader(input_files=[file_path]).load_data()
                else:
                    # PDF Document Handling (Optimized via pypdfium2)
                    if file_path.lower().endswith('.pdf'):
                        try:
                            import pypdfium2 as pdfium
                            import threading
                            if not hasattr(pdfium, "_global_thread_lock"):
                                pdfium._global_thread_lock = threading.Lock()
                                
                            print(f"[DocumentIngestor] Extracting PDF text via fast pypdfium2...")
                            
                            documents = []
                            with pdfium._global_thread_lock:
                                with pdfium.PdfDocument(file_path) as pdf:
                                    for i, page in enumerate(pdf):
                                        textpage = page.get_textpage()
                                        page_text = textpage.get_text_bounded()
                                        
                                        # Create a LlamaIndex Document for each page
                                        doc = Document(
                                            text=page_text or "",
                                            metadata={
                                                "page_label": str(i + 1),
                                                "file_name": os.path.basename(file_path)
                                            }
                                        )
                                        documents.append(doc)
                            
                            print(f"[DocumentIngestor] Fast PDF extraction successful: {len(documents)} pages extracted.")
                        except Exception as e:
                            print(f"[DocumentIngestor] Fast PDF extraction failed, falling back: {e}")
                            documents = SimpleDirectoryReader(input_files=[file_path]).load_data()
                    else:
                        documents = SimpleDirectoryReader(input_files=[file_path]).load_data()
                
                # Metadata sanitization
                if documents:
                    for doc in documents:
                        for key in list(doc.metadata.keys()):
                            val = doc.metadata[key]
                            if isinstance(val, (list, dict, set, tuple)) or val is None:
                                 del doc.metadata[key]
                                 continue
                            if isinstance(val, str) and len(val) > 400:
                                 doc.metadata[key] = val[:400] + "..."

            # Processing documents
            if documents:
                for doc in documents:
                    if metadata:
                        doc.metadata.update(metadata)
                    doc.metadata["source"] = file_path
                    doc.metadata["type"] = "document"
                    if conversation_id:
                        doc.metadata["conversation_id"] = conversation_id
                    
                    doc.excluded_llm_metadata_keys = ["conversation_id", "source"]
                    doc.excluded_embed_metadata_keys = ["conversation_id", "source"]

                # Split into nodes
                from llama_index.core import Settings
                nodes = []
                enable_metadata = metadata.get("enable_metadata", False) if metadata else False
                
                if enable_metadata:
                    try:
                        from llama_index.core.ingestion import IngestionPipeline
                        from llama_index.core.extractors import TitleExtractor, SummaryExtractor
                        active_llm = llm or Settings.llm
                        
                        if not active_llm:
                             nodes = Settings.node_parser.get_nodes_from_documents(documents)
                        else:
                            pipeline = IngestionPipeline(
                                transformations=[
                                    Settings.node_parser,
                                    TitleExtractor(nodes=5, llm=active_llm),
                                    SummaryExtractor(summaries=["prev", "self", "next"], llm=active_llm),
                                ]
                            )
                            nodes = pipeline.run(documents=documents)
                    except Exception as e:
                        print(f"[DocumentIngestor] Error in extraction pipeline: {e}")
                        nodes = Settings.node_parser.get_nodes_from_documents(documents)
                else:
                    nodes = Settings.node_parser.get_nodes_from_documents(documents)

                # Prepare metadata and exclusions
                total_nodes = len(nodes)
                for node in nodes:
                    if "_node_content" not in node.excluded_embed_metadata_keys:
                        node.excluded_embed_metadata_keys.append("_node_content")
                    if "_node_content" not in node.excluded_llm_metadata_keys:
                        node.excluded_llm_metadata_keys.append("_node_content")

                # Insert nodes in batches of 100 to optimize API calls while giving progress updates
                batch_size = 100
                for i in range(0, total_nodes, batch_size):
                    batch = nodes[i:i + batch_size]
                    index.insert_nodes(batch)
                    if progress_callback:
                        progress_callback(min(i + batch_size, total_nodes), total_nodes)
                
                # Result construction
                text = "\n\n".join([doc.text for doc in documents])
                print(f"[DocumentIngestor] Ingestion complete. Text length: {len(text)}. Nodes: {len(nodes)}")
                
                return {
                    "status": "success",
                    "count": len(nodes),
                    "full_text": text,
                    "text_length": len(text),
                    "doc_count": len(documents),
                    "structured_data": structured_data
                }
            
            return {"status": "no_content", "full_text": ""}

        except Exception as e:
            import traceback
            print(f"[DocumentIngestor] CRITICAL ERROR: {e}")
            traceback.print_exc()
            return {"status": "error", "error": str(e)}

document_ingestor = DocumentIngestor()
